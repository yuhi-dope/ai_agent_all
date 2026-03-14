"""File ingestion — PDF, Excel, CSV, images via OCR.

MVP: text extraction from common file types.
Phase 2+: Google Document AI for handwritten/complex documents.
"""
import csv
import io
import logging
from uuid import UUID

from brain.extraction import ExtractionResult, extract_knowledge
from db.supabase import get_service_client

logger = logging.getLogger(__name__)

SUPPORTED_TEXT_TYPES = {
    "text/plain", "text/csv",
    "application/json",
}

SUPPORTED_DOC_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",  # xlsx
    "application/vnd.ms-excel",  # xls
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # docx
    "application/msword",  # doc
    "image/jpeg", "image/png", "image/webp",
}


async def ingest_file(
    file_content: bytes,
    filename: str,
    content_type: str,
    company_id: str,
    user_id: str,
    department: str | None = None,
    category: str | None = None,
) -> ExtractionResult:
    """Ingest a file by extracting text and running extraction pipeline.

    Supports: .txt, .csv, .json, .pdf, .xlsx, .xls, .docx.
    Phase 2+: Google Document AI for images/handwritten docs.
    """
    text = await _extract_text(file_content, filename, content_type)

    if not text or not text.strip():
        raise ValueError(f"Could not extract text from {filename} ({content_type})")

    # Run extraction pipeline (creates session + saves items internally)
    result = await extract_knowledge(
        text=text,
        company_id=company_id,
        user_id=user_id,
        department=department,
        category=category,
    )

    return result


async def _extract_text(content: bytes, filename: str, content_type: str) -> str:
    """Extract readable text from file content."""
    # Plain text
    if content_type in ("text/plain", "application/json"):
        return content.decode("utf-8", errors="replace")

    # CSV
    if content_type == "text/csv" or filename.endswith(".csv"):
        return _extract_csv(content)

    # PDF (basic text extraction)
    if content_type == "application/pdf" or filename.endswith(".pdf"):
        return await _extract_pdf(content)

    # Word documents (.docx)
    if content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or filename.endswith((".docx", ".doc")):
        return await _extract_docx(content, filename)

    # Images — Phase 2+ (Document AI)
    if content_type.startswith("image/"):
        raise ValueError(
            f"Image OCR is not yet supported (Phase 2+). "
            f"Please convert to text and use text ingestion."
        )

    # Excel — Phase 1+ (requires openpyxl)
    if "spreadsheet" in content_type or filename.endswith((".xlsx", ".xls")):
        return await _extract_excel(content, filename)

    raise ValueError(f"Unsupported file type: {content_type} ({filename})")


def _extract_csv(content: bytes) -> str:
    """Convert CSV to readable text."""
    text = content.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)

    if not rows:
        return ""

    # Format as readable text with headers
    headers = rows[0]
    lines = []
    for row in rows[1:]:
        parts = [f"{h}: {v}" for h, v in zip(headers, row) if v.strip()]
        if parts:
            lines.append("、".join(parts))

    return "\n".join(lines)


async def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF using PyPDF2 (basic).

    Phase 2+: Replace with Google Document AI for complex/scanned PDFs.
    """
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        texts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
        return "\n\n".join(texts)
    except ImportError:
        raise ValueError(
            "PDF extraction requires pypdf. Install with: pip install pypdf"
        )
    except Exception as e:
        raise ValueError(f"PDF extraction failed: {e}")


async def _extract_excel(content: bytes, filename: str) -> str:
    """Extract text from Excel using openpyxl (basic)."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True)
        texts = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            headers = [str(h) if h else f"列{i}" for i, h in enumerate(rows[0])]
            for row in rows[1:]:
                parts = [f"{h}: {v}" for h, v in zip(headers, row) if v is not None]
                if parts:
                    texts.append("、".join(parts))
        wb.close()
        return "\n".join(texts)
    except ImportError:
        raise ValueError(
            "Excel extraction requires openpyxl. Install with: pip install openpyxl"
        )
    except Exception as e:
        raise ValueError(f"Excel extraction failed: {e}")


async def _extract_docx(content: bytes, filename: str) -> str:
    """Extract text from Word documents using python-docx.

    Supports .docx files. For legacy .doc files, python-docx may not work;
    Phase 2+ will add full .doc support via LibreOffice conversion.
    """
    if filename.endswith(".doc") and not filename.endswith(".docx"):
        raise ValueError(
            "Legacy .doc format is not fully supported yet. "
            "Please convert to .docx and retry."
        )
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        texts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                texts.append(text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    texts.append("、".join(cells))

        return "\n".join(texts)
    except ImportError:
        raise ValueError(
            "Word extraction requires python-docx. Install with: pip install python-docx"
        )
    except Exception as e:
        raise ValueError(f"Word extraction failed: {e}")
